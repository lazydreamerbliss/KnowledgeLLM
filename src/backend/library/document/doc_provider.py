import os
from datetime import datetime
from time import time
from typing import Any, Generator

import docx
from docx.document import Document
from library.document.doc_lib_table import DocLibTable
from library.document.doc_provider_base import *
from loggers import doc_lib_logger as LOGGER
from pypdf import PdfReader
from utils.task_runner import report_progress

TXT_EXTENSION: str = 'txt'
PDF_EXTENSION: str = 'pdf'
DOC_EXTENSION: str = 'doc'
DOCX_EXTENSION: str = 'docx'
EBOOK_EXTENSIONS: set[str] = {'epub', 'mobi', 'azw', 'azw3', 'azw4', 'prc', 'tpz'}


class DocProvider(DocProviderBase[DocLibTable]):
    """A generic type of DocLibTable as the type of SQL table is needed
    """

    TABLE_TYPE: type = DocLibTable
    DOC_TYPE: str = DocumentType.GENERAL.value

    def __init__(self,
                 db_path: str,
                 uuid: str,
                 doc_path: str | None = None,
                 progress_reporter: Callable[[int, int, str | None], None] | None = None):
        super().__init__(db_path, uuid, doc_path, progress_reporter, table_type=DocProvider.TABLE_TYPE)

        # If the table is empty, initialize it with given doc_path
        if not self._table.row_count():
            if not doc_path:
                raise DocProviderError('doc_path is mandatory when table is empty')
            LOGGER.info(f'Document table is empty, initializing document: {doc_path}...')
            self.initialize(doc_path)

    def __reader_pdf(self, doc_path: str) -> Generator[tuple[str, int], Any, None]:
        """Reader for .pdf document

        Yields:
            Generator[tuple[str, int], Any, None]: line, current progress (%)
        """
        try:
            reader: PdfReader = PdfReader(doc_path)
            total: int = len(reader.pages)
            for i, page in enumerate(reader.pages):
                page_text: str = page.extract_text()
                page_text = page_text.strip()
                if not page_text:
                    continue

                # PDF text is separated by '\n \n', single '\n' is for line wrapping
                for line in page_text.split('\n \n'):
                    line = line.strip()
                    if not line:
                        continue
                    line = ''.join([s.strip() for s in line.split('\n')])
                    yield line, int(i / total * 100)
            return
        except Exception as e:
            msg: str = f'Failed to read PDF content: {doc_path}, error: {e}'
            LOGGER.error(msg)
            raise DocProviderError(msg)

    def __reader_docx(self, doc_path: str) -> Generator[tuple[str, int], Any, None]:
        """Reader for .docx document

        Yields:
            Generator[tuple[str, int], Any, None]: line and total number of lines
        """
        try:
            doc: Document = docx.Document(doc_path)
            total: int = len(doc.paragraphs)
            for i, paragraph in enumerate(doc.paragraphs):
                line = paragraph.text.strip()
                if not line:
                    continue
                yield line, int(i / total * 100)
            return
        except Exception as e:
            msg: str = f'Failed to read DOC/DOCX content: {doc_path}, error: {e}'
            LOGGER.error(msg)
            raise DocProviderError(msg)

    def __reader_ebook(self, doc_path: str) -> Generator[tuple[str, int], Any, None]:
        """Reader for EBook document

        Yields:
            Generator[tuple[str, int], Any, None]: line and total number of lines
        """
        raise NotImplementedError()

    def __reader_txt(self, doc_path: str) -> Generator[tuple[str, int], Any, None]:
        """Reader for text file document

        Yields:
            Generator[tuple[str, int], Any, None]: line and total number of lines
        """
        try:
            with open(doc_path, 'r', encoding='utf-8') as f:
                all_lines: list[str] = f.readlines()
            total: int = len(all_lines)
            for i, line in enumerate(all_lines):
                line = line.strip()
                if not line:
                    continue
                yield line, int(i / total * 100)
        except Exception as e:
            msg: str = f'Failed to read plain text content: {doc_path}, error: {e}'
            LOGGER.error(msg)
            raise DocProviderError(msg)

    def initialize(self, doc_path: str) -> None:
        if not doc_path:
            raise DocProviderError('doc_path is None')

        _, extension = os.path.splitext(doc_path)
        if not extension:
            msg: str = f'Unknown document type: {doc_path}'
            LOGGER.error(msg)
            raise DocProviderError(msg)
        extension = extension.lower()[1:]
        start: float = time()

        LOGGER.info(f'Reading document: {doc_path}...')
        reader: Callable[[str], Generator[tuple[str, int], Any, None]] | None = None
        if extension == PDF_EXTENSION:
            reader = self.__reader_pdf
        elif extension == DOC_EXTENSION or extension == DOCX_EXTENSION:
            reader = self.__reader_docx
        elif extension in EBOOK_EXTENSIONS:
            reader = self.__reader_ebook
        elif extension == TXT_EXTENSION:
            reader = self.__reader_txt
        if not reader:
            raise DocProviderError(f'Unsupported document format: {extension}')

        timestamp: datetime = datetime.now()
        previous_progress: int = -1
        i: int = 0
        for line, current_progress in reader(doc_path):
            i += 1
            if not line:
                continue

            # If reporter is given, report progress to task manager
            # - Reduce report frequency, only report when progress changes
            if current_progress > previous_progress:
                previous_progress = current_progress
                report_progress(self._progress_reporter, current_progress, current_phase=1, phase_name='DUMP')

            self._table.insert_row((timestamp, line))

        time_taken: float = time() - start
        LOGGER.info(f'Finished processing document: {doc_path}, cost: {time_taken:.2f}s')

    def get_key_text_from_record(self, row: tuple) -> str:
        # The text column ['text', 'TEXT'] is the 3rd column of document table, so row[2] is the key info
        return row[2]
